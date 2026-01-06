import docx

def create_sample_docx(filename):
    doc = docx.Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('This is a sample document for testing text extraction.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This section contains some logic about how to process files.')
    doc.add_paragraph('1. Read file\n2. Extract text\n3. Generate reasoning')
    doc.save(filename)

if __name__ == "__main__":
    create_sample_docx("my_workspace/sample_doc.docx")
    print("Created my_workspace/sample_doc.docx")
